import json
import os
from datetime import datetime
from io import StringIO
from typing import List

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell

from config import Config
from logfile import Logging as l


class DocUtils:
    out_directory = "test_outputs" if Config.testing else "outputs"

    @staticmethod
    def word_table_to_df(doc_path: str, table_index=0) -> pd.DataFrame:
        doc: Document = Document(doc_path)
        table: Table = doc.tables[table_index]

        data: List[List[str]] = []
        keys = None
        for i, row in enumerate(table.rows):
            text = [cell.text for cell in row.cells]
            if i == 0:
                keys = text
            else:
                data.append(text)

        df = pd.DataFrame(data, columns=keys)
        return df

    @staticmethod
    def df_to_json_string(df: pd.DataFrame) -> str:
        return json.dumps(json.loads((df.to_json())), indent=4)

    @staticmethod
    def json_string_to_df(json_string: str) -> pd.DataFrame:
        return pd.read_json(StringIO(json_string))

    @staticmethod
    def get_total_hours(df: pd.DataFrame) -> float:
        df['Hours'] = pd.to_numeric(df['Hours'], errors='coerce')
        return df['Hours'].sum()

    @staticmethod
    def df_to_doc(df: pd.DataFrame) -> Document:
        df = df.fillna('')
        doc: Document = Document()
        doc.add_heading('Apprenticeship Log', level=1)

        table: Table = doc.add_table(rows=1, cols=len(df.columns))

        hdr_cells = table.rows[0].cells
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = column
            DocUtils.set_cell_border(hdr_cells[i])

        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)
                DocUtils.set_cell_border(row_cells[i])

        return doc

    @staticmethod
    def set_cell_border(cell: _Cell, border_type: str = "single", size: int = 6, color: str = "000000",
                        space: int = 0) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        borders = tcPr.find(qn('w:tcBorders'))
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tcPr.append(borders)

        for edge in ('top', 'start', 'bottom', 'end'):
            edge_tag = f'w:{edge}'
            element = OxmlElement(edge_tag)
            element.set(qn('w:val'), border_type)
            element.set(qn('w:sz'), str(size))
            element.set(qn('w:space'), str(space))
            element.set(qn('w:color'), color)
            borders.append(element)

    @staticmethod
    def save_doc(doc: Document, week: str, hours: float) -> str:
        filename = DocUtils.create_filename(hours, Config.name, week)
        directory: str = DocUtils.get_directory(week)
        full_path: str = f"../{DocUtils.out_directory}/{directory}/{filename}"
        doc.save(full_path)
        l.log_with_timestamp(f"Document {full_path} saved.")
        return filename

    @staticmethod
    def create_filename(hours, name, date_str) -> str:
        date = datetime.strptime(date_str, "%d/%m/%Y")
        formatted_date = date.strftime("%d-%m-%Y")
        safe_name = name.replace(" ", "_")
        return f"ActivityLog_{safe_name}_{hours}_{formatted_date}.docx"

    @staticmethod
    def get_directory(date: str) -> str:
        date_obj = datetime.strptime(date, "%d/%m/%Y")
        month, year = date_obj.strftime("%B"), date_obj.strftime("%Y")
        month_directory: str = DocUtils.make_directory_name(month, year)
        if not os.path.exists(f'../{DocUtils.out_directory}/{month_directory}'):
            os.makedirs(f'../{DocUtils.out_directory}/{month_directory}')
        return month_directory

    @staticmethod
    def make_directory_name(month: str, year: str) -> str:
        month_order = {
            "January": "01", "February": "02", "March": "03",
            "April": "04", "May": "05", "June": "06",
            "July": "07", "August": "08", "September": "09",
            "October": "10", "November": "11", "December": "12"
        }
        return f"{month_order[month]}_{month}-{year}"
